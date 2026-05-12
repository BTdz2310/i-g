#!/usr/bin/env python3
"""Generate partner integration guide as a DOCX file."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------- Styles ----------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2')
    rpr.append(shd)
    return p

def add_h(level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h

def add_p(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p

def add_kv_table(rows, col_widths=(Inches(1.8), Inches(4.5))):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(rows):
        c1 = t.rows[i].cells[0]
        c2 = t.rows[i].cells[1]
        c1.width = col_widths[0]
        c2.width = col_widths[1]
        c1.text = ''
        c2.text = ''
        r1 = c1.paragraphs[0].add_run(k)
        r1.bold = True
        c2.paragraphs[0].add_run(v)
        set_cell_bg(c1, 'F2F4F8')
    return t

# ---------- Cover ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('TÀI LIỆU TÍCH HỢP\nINSURANCE GATEWAY (PVIS)')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run('Dành cho đối tác cấp 2 — Bảo hiểm TNDS xe cơ giới')
sr.italic = True
sr.font.size = Pt(13)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run('Phiên bản 1.0  •  Tháng 05/2026  •  PVI Phía Nam')
mr.font.size = Pt(10)
mr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

# ---------- 1. Thông tin chung ----------
add_h(1, '1. Thông tin chung')
add_p(
    'Insurance Gateway là hệ thống trung gian do PVIS vận hành, đứng giữa hệ thống của đối tác '
    'và PVI Core (Tổng công ty Bảo hiểm PVI). Đối tác chỉ cần tích hợp một lần với gateway này — '
    'không phải tích hợp trực tiếp với PVI Core.'
)
add_p('Nhiệm vụ chính của gateway:')
for x in [
    'Xác thực đối tác bằng HMAC-SHA256 (mạnh hơn MD5 mà PVI Core dùng).',
    'Tạo / tra cứu / đối soát đơn bảo hiểm TNDS xe cơ giới.',
    'Nhận callback từ PVI và lưu kết quả cấp đơn.',
    'Cung cấp endpoint tra cứu giao dịch theo đối tác.',
]:
    doc.add_paragraph(x, style='List Bullet')

# ---------- 2. Thông tin tài khoản ----------
add_h(1, '2. Thông tin tài khoản (PVIS cấp riêng cho từng đối tác)')
add_p(
    'Các giá trị dưới đây do PVIS cấp riêng cho mỗi đối tác. Tuyệt đối không chia sẻ secret cho '
    'bên thứ ba. Khi nghi ngờ bị lộ, liên hệ ngay PVIS để rotate.'
)

add_h(2, '2.1. Endpoint & môi trường')
add_kv_table([
    ('Base URL (UAT)', '_______________________________________'),
    ('Base URL (Production)', '_______________________________________'),
])

add_h(2, '2.2. Credential HMAC partner')
add_kv_table([
    ('Client ID', '_______________________________________'),
    ('Key ID', '_______________________________________'),
    ('Secret (hex 64 ký tự)', '_______________________________________'),
    ('IP allowlist (nếu có)', '_______________________________________'),
    ('Rate limit (req/phút)', '_______________________________________'),
])
add_p(
    'Lưu ý: Secret chỉ được PVIS cấp 1 lần khi tạo partner hoặc rotate. Không thể xem lại sau đó. '
    'Nếu mất, phải yêu cầu rotate-secret và tích hợp lại.'
)

add_h(2, '2.3. Tài khoản admin (chỉ cho người quản trị bên đối tác, nếu được cấp)')
add_kv_table([
    ('Admin username', '_______________________________________'),
    ('Admin password', '_______________________________________'),
])
add_p(
    'Tài khoản admin dùng để đăng nhập trang quản trị (xem báo cáo, danh sách giao dịch...). '
    'Tài khoản này KHÔNG dùng để gọi API tích hợp — API tích hợp chỉ dùng credential HMAC ở mục 2.2.'
)

# ---------- 3. Xác thực HMAC ----------
add_h(1, '3. Cơ chế xác thực HMAC (6 headers)')
add_p(
    'Mỗi request gọi đến gateway (trừ /admin/auth/login và /pvi/callback) phải đính kèm 6 header '
    'dưới đây. Gateway kiểm tra: chữ ký HMAC, thời gian lệch ≤ 5 phút, nonce duy nhất ≤ 5 phút, '
    'và rate limit theo phút.'
)

add_h(2, '3.1. Danh sách headers')
hdr_rows = [
    ('Header', 'Giá trị', 'Cách tính'),
    ('X-Client-Id', 'Client ID được cấp', 'Cố định, lấy từ mục 2.2'),
    ('X-Key-Id', 'Key ID được cấp', 'Cố định, lấy từ mục 2.2'),
    ('X-Timestamp', 'Unix epoch giây', 'Math.floor(Date.now() / 1000) — sai lệch ≤ 5 phút so với server'),
    ('X-Nonce', 'UUID v4 ngẫu nhiên', 'crypto.randomUUID() — mỗi request 1 giá trị, không tái sử dụng trong 5 phút'),
    ('X-Signature-Version', 'v1', 'Cố định chuỗi "v1"'),
    ('X-Signature', 'HMAC-SHA256 hex (64 ký tự)', 'Xem mục 3.2 — tính từ secret + canonical string'),
]
t = doc.add_table(rows=len(hdr_rows), cols=3)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(hdr_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

# ---------- 3.2 Cách tính X-Signature ----------
add_h(2, '3.2. Cách tính X-Signature (khó nhất, đọc kỹ)')
add_p('Công thức tổng quát:')
add_code(
    'X-Signature = hex( HMAC-SHA256( secret, canonical_string ) )'
)

add_p('Trong đó canonical_string là 5 thành phần nối bằng ký tự xuống dòng \\n:', bold=True)
add_code(
    'canonical_string =\n'
    '    METHOD                    + "\\n" +\n'
    '    PATH_WITH_QUERY           + "\\n" +\n'
    '    TIMESTAMP                 + "\\n" +\n'
    '    NONCE                     + "\\n" +\n'
    '    BODY_SHA256_HEX'
)

add_p('Quy tắc từng thành phần:', bold=True)
rules = [
    ('METHOD', 'HTTP method viết HOA. Ví dụ: "POST", "GET", "PATCH". Không phải "Post".'),
    ('PATH_WITH_QUERY', 'Path + query string nguyên văn như đã gửi đi. KHÔNG sort lại query. Ví dụ: "/api/pvi/catalog" hoặc "/transaction?status=ISSUED&from=2026-01-01".'),
    ('TIMESTAMP', 'Lấy đúng giá trị header X-Timestamp (ví dụ "1747300000"). Không cộng/trừ.'),
    ('NONCE', 'Lấy đúng giá trị header X-Nonce.'),
    ('BODY_SHA256_HEX', 'SHA-256 của raw request body (byte), encode hex 64 ký tự lowercase. Nếu request KHÔNG có body (GET, DELETE...) thì dùng SHA-256 của chuỗi rỗng "".'),
]
for k, v in rules:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(k + ': ')
    r1.bold = True
    p.add_run(v)

add_p('Lưu ý cực kỳ quan trọng về BODY:', bold=True)
caveats = [
    'Body dùng để hash phải là CHÍNH XÁC chuỗi bytes sẽ gửi đi. Nếu bạn JSON.stringify rồi gửi, hash phải tính trên kết quả stringify đó.',
    'Nếu bạn format / pretty-print / sắp xếp lại key sau khi hash → server sẽ tính ra hash khác → 401.',
    'Cách an toàn: stringify đúng 1 lần, lưu vào biến, hash biến đó, gửi biến đó làm body.',
    'SHA-256 của body rỗng (cho GET) luôn là: e3b0c44298fc1c149afbf4c8996fb92427ae41e4649b934ca495991b7852b855',
]
for c in caveats:
    doc.add_paragraph(c, style='List Bullet')

# ---------- 3.3 Ví dụ tính tay ----------
add_h(2, '3.3. Ví dụ tính tay (debug cho dễ)')
add_p('Giả sử bạn có:')
add_kv_table([
    ('Secret', '9e4e91c78df3b1c36b3e212484c8df7d6e7a953a10e062a2107311eaf7263675'),
    ('Method', 'POST'),
    ('Path', '/api/pvi/catalog'),
    ('Timestamp', '1747300000'),
    ('Nonce', 'a4d8c7b2-1234-4abc-9def-0123456789ab'),
    ('Body (đã stringify)', '{"ten_dmuc":"LOAIXEMOTOR","parent_value":"1","giatri_chon":""}'),
])

add_p('Bước 1 — SHA-256 của body (hex lowercase):')
add_code('BODY_SHA256_HEX = sha256(\'{"ten_dmuc":"LOAIXEMOTOR","parent_value":"1","giatri_chon":""}\')')

add_p('Bước 2 — Ghép canonical_string (chú ý có \\n giữa các dòng, không có \\n cuối):')
add_code(
    'POST\n'
    '/api/pvi/catalog\n'
    '1747300000\n'
    'a4d8c7b2-1234-4abc-9def-0123456789ab\n'
    '<BODY_SHA256_HEX vừa tính ở bước 1>'
)

add_p('Bước 3 — HMAC-SHA256(secret, canonical_string) rồi encode hex → đó là X-Signature.')

# ---------- 3.4 Code mẫu ----------
add_h(2, '3.4. Code mẫu sinh chữ ký')

add_p('Node.js / TypeScript:', bold=True)
add_code('''import crypto from 'crypto';

function signRequest(opts: {
  secret: string;
  clientId: string;
  keyId: string;
  method: string;
  pathWithQuery: string;
  body: string; // chuỗi đã stringify, hoặc "" nếu không có body
}) {
  const timestamp = Math.floor(Date.now() / 1000).toString();
  const nonce = crypto.randomUUID();
  const bodyHash = crypto.createHash('sha256').update(opts.body, 'utf8').digest('hex');

  const canonical = [
    opts.method.toUpperCase(),
    opts.pathWithQuery,
    timestamp,
    nonce,
    bodyHash,
  ].join('\\n');

  const signature = crypto.createHmac('sha256', opts.secret)
    .update(canonical, 'utf8')
    .digest('hex');

  return {
    'X-Client-Id': opts.clientId,
    'X-Key-Id': opts.keyId,
    'X-Timestamp': timestamp,
    'X-Nonce': nonce,
    'X-Signature-Version': 'v1',
    'X-Signature': signature,
  };
}

// Sử dụng:
const body = JSON.stringify({ ten_dmuc: 'LOAIXEMOTOR', parent_value: '1', giatri_chon: '' });
const headers = signRequest({
  secret: 'YOUR_SECRET',
  clientId: 'YOUR_CLIENT_ID',
  keyId: 'YOUR_KEY_ID',
  method: 'POST',
  pathWithQuery: '/api/pvi/catalog',
  body,
});
await fetch('https://gateway.pvis.vn/api/pvi/catalog', {
  method: 'POST',
  headers: { 'Content-Type': 'application/json', ...headers },
  body, // PHẢI là cùng chuỗi đã hash
});''')

add_p('Python 3:', bold=True)
add_code('''import hmac, hashlib, time, uuid, json, requests

def sign_request(secret, client_id, key_id, method, path_with_query, body=''):
    timestamp = str(int(time.time()))
    nonce = str(uuid.uuid4())
    body_hash = hashlib.sha256(body.encode('utf-8')).hexdigest()
    canonical = '\\n'.join([method.upper(), path_with_query, timestamp, nonce, body_hash])
    signature = hmac.new(secret.encode('utf-8'), canonical.encode('utf-8'), hashlib.sha256).hexdigest()
    return {
        'X-Client-Id': client_id,
        'X-Key-Id': key_id,
        'X-Timestamp': timestamp,
        'X-Nonce': nonce,
        'X-Signature-Version': 'v1',
        'X-Signature': signature,
    }

body = json.dumps({'ten_dmuc': 'LOAIXEMOTOR', 'parent_value': '1', 'giatri_chon': ''}, separators=(',', ':'))
headers = sign_request('SECRET', 'CLIENT_ID', 'KEY_ID', 'POST', '/api/pvi/catalog', body)
r = requests.post('https://gateway.pvis.vn/api/pvi/catalog',
                  headers={'Content-Type': 'application/json', **headers},
                  data=body)  # data=body để gửi đúng chuỗi đã hash
print(r.status_code, r.json())''')

add_p('Phản hồi khi xác thực fail:', bold=True)
fail_rows = [
    ('Status', 'Nguyên nhân thường gặp'),
    ('400 Missing authentication headers', 'Thiếu 1 trong 6 header'),
    ('401 Invalid signature version', 'X-Signature-Version khác "v1"'),
    ('401 Invalid credentials', 'Client ID / Key ID không đúng hoặc đã bị disable'),
    ('401 Stale timestamp', 'Thời gian client lệch quá 5 phút so với server'),
    ('401 Invalid signature', 'Tính chữ ký sai — kiểm tra lại body bytes & canonical string'),
    ('401 Replay detected', 'Nonce đã dùng — sinh nonce mới cho mỗi request'),
    ('403 IP not allowed', 'IP nguồn không nằm trong allowlist của partner'),
    ('429 Rate limit exceeded', 'Vượt quá rate limit/phút'),
]
t = doc.add_table(rows=len(fail_rows), cols=2)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(fail_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

# ---------- 4. Endpoints ----------
add_h(1, '4. Danh sách endpoints')

# 4.1 Catalog
add_h(2, '4.1. POST /api/pvi/catalog — Lấy danh mục')
add_p('Dùng để lấy các danh sách dropdown: loại xe, hãng xe, dòng xe, mục đích sử dụng...')
add_p('Body:')
add_code('''{
  "ten_dmuc": "<tên danh mục>",
  "parent_value": "<value của mục cha, để '' nếu không có>",
  "giatri_chon": "<giá trị mặc định đã chọn, để '' nếu không có>"
}''')

add_p('Các giá trị ten_dmuc thường dùng:')
cat_rows = [
    ('ten_dmuc', 'Mô tả', 'parent_value'),
    ('LOAIXEMOTOR', 'Loại xe máy', 'Bắt buộc = "1"'),
    ('HIEUXEMOTOR', 'Hiệu xe máy', '""'),
    ('LOAIXEAUTO', 'Loại xe ô tô', '""'),
    ('HIEUXEAUTO', 'Hiệu xe ô tô', '""'),
    ('DONGXE', 'Dòng xe (sau khi có hiệu xe)', 'value của hiệu xe'),
    ('MDSD_AUTO', 'Mục đích sử dụng xe ô tô', '""'),
    ('LOAIHINH_AUTO', 'Loại hình theo mục đích', 'value của mục đích'),
]
t = doc.add_table(rows=len(cat_rows), cols=3)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(cat_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

add_p('Ví dụ cụ thể — lấy danh sách LOẠI XE MÁY:', bold=True)
add_code('''POST /api/pvi/catalog
{
  "ten_dmuc": "LOAIXEMOTOR",
  "parent_value": "1",
  "giatri_chon": ""
}''')
add_p(
    'Lưu ý: với LOAIXEMOTOR thì parent_value BẮT BUỘC bằng "1". Truyền giá trị khác sẽ không trả về danh sách đúng.'
)

add_p('Response:')
add_code('''[
  { "Value": "20011", "Text": "Xe máy dưới 50cc" },
  { "Value": "20012", "Text": "Xe máy 50cc trở lên" },
  ...
]''')

# 4.2 Vehicle type
add_h(2, '4.2. POST /api/pvi/vehicle-type — Lấy mã loại xe ô tô')
add_p('Dùng để lấy ma_loaixe cho ô tô dựa trên số chỗ + trọng tải + mục đích sử dụng. Kết quả dùng tiếp ở /quote.')
add_code('''POST /api/pvi/vehicle-type
{
  "SoChoNgoi": 5,
  "TrongTai": 0,
  "Ma_MDSD": "1",
  "LoaiHinh": ""
}''')

# 4.3 Quote
add_h(2, '4.3. POST /api/pvi/quote — Tính phí TNDS')
add_p(
    'Tính tổng phí bảo hiểm. Phải gọi trước khi tạo đơn để biết PhiBHTNDSBB và TongPhi.'
)

add_h(3, '4.3.1. Cách tính thời hạn bảo hiểm')
add_p(
    'Đối tác chỉ cần cho khách chọn THỜI HẠN theo NĂM (1 / 2 / 3 năm). Hệ thống đối tác tự tính ngaycuoi từ ngaydau theo công thức:'
)
add_code('ngaycuoi = ngaydau + số_năm  (cùng ngày/tháng, năm + N)')
add_p('Ví dụ:')
ex_rows = [
    ('ngaydau', 'Thời hạn chọn', 'ngaycuoi'),
    ('01/06/2026', '1 năm', '01/06/2027'),
    ('01/06/2026', '2 năm', '01/06/2028'),
    ('15/03/2026', '3 năm', '15/03/2029'),
]
t = doc.add_table(rows=len(ex_rows), cols=3)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(ex_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

add_h(3, '4.3.2. Bảo hiểm trách nhiệm bồi thường người ngồi trên xe (BH lái phụ / hành khách)')
add_p(
    'Nếu khách muốn mua kèm gói "Bảo hiểm trách nhiệm bồi thường đối với người ngồi trên xe", '
    'truyền 2 trường sau khi gọi /quote:'
)
qrow = [
    ('Trường', 'Giá trị', 'Ghi chú'),
    ('thamgia_laiphu', 'true', 'Bật gói BH lái phụ'),
    ('mtn_laiphu', 'Số tiền bảo hiểm (VND)', 'Mức trách nhiệm khách hàng chọn, ví dụ 10000000, 20000000...'),
]
t = doc.add_table(rows=len(qrow), cols=3)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(qrow):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

add_p('Nếu KHÔNG mua kèm: gửi thamgia_laiphu = false và mtn_laiphu = 0 (hoặc bỏ trống).')

add_h(3, '4.3.3. Body đầy đủ của /quote')
add_code('''POST /api/pvi/quote
{
  "ma_loaixe": "20012",
  "ma_trongtai": "",
  "so_cho": 0,
  "ma_mdsd": "1",
  "giodau": "00:00",
  "giocuoi": "23:59",
  "ngaydau": "01/06/2026",
  "ngaycuoi": "01/06/2027",
  "thamgia_tndsbb": true,
  "thamgia_laiphu": true,
  "mtn_laiphu": 10000000,
  "so_nguoi": 2,
  "philpx_nhap": 0
}''')

add_p('Response:')
add_code('''{
  "Status": "00",
  "Message": "Thanh cong",
  "TotalFee": "510700",
  "phi_tndsbb": "480700",
  "phi_lpx": "30000",
  "ma_loaixe": "20012"
}''')
add_p('TotalFee là số tiền cuối cùng để hiển thị cho khách và truyền lại vào /order ở trường TongPhi.')

# 4.4 Create order
add_h(2, '4.4. POST /api/pvi/order — Tạo đơn bảo hiểm')
add_p(
    'Gateway tự sinh maGiaodich (UUID) và trả về. Đơn sẽ được PVI cấp bất đồng bộ — '
    'kết quả về qua callback hoặc qua GET /api/pvi/order/:maGiaodich.'
)
add_code('''POST /api/pvi/order
{
  "TenKH": "Nguyễn Văn A",
  "DiaChiKH": "123 Lê Lợi, Q.1, TP.HCM",
  "TenChuXe": "Nguyễn Văn A",
  "DiaChiChuXe": "123 Lê Lợi, Q.1, TP.HCM",
  "EmailKH": "khachhang@email.com",
  "DienThoai": "0912345678",

  "NgayDau": "01/06/2026",
  "NgayCuoi": "01/06/2027",
  "GioDau": "00:00",
  "GioCuoi": "23:59",

  "LoaiXe": "20012",
  "TenLoaiXe": "Xe máy 50cc trở lên",
  "ChoNgoi": "2",
  "TrongTai": "0",
  "HieuXe": "...",
  "DongXe": "...",
  "BienKiemSoat": "59X1-12345",
  "SoKhung": "VIN...",
  "SoMay": "ENG...",
  "NamSX": "01/2024",
  "NamSD": "01/2024",
  "MaMucDichSD": "1",

  "PhiBHTNDSBB": "480700",
  "TongPhi": "510700",

  "ThamGiaLaiPhu": true,
  "MTNLaiPhu": "10000000",
  "SoNguoiToiDa": "2",
  "PhiBHLaiPhu": "30000",

  "productKind": "MOTO"
}''')
add_p('Response:')
add_code('''{
  "maGiaodich": "9e8d3c44-...-...-...",
  "Pr_key": 123456,
  "paymentUrl": "https://piastest.pvi.com.vn/payment/...",
  "serialNumber": null
}''')
add_p('Quy trình tiếp theo: redirect khách sang paymentUrl. Sau khi thanh toán + PVI cấp đơn xong, callback sẽ về gateway và đối tác có thể GET /api/pvi/order/:maGiaodich để lấy GCN PDF.')

# 4.5 Get policy
add_h(2, '4.5. GET /api/pvi/order/:maGiaodich — Tra cứu đơn')
add_p('Trả trạng thái hiện tại của đơn. Đối tác có thể polling endpoint này.')
add_code('''GET /api/pvi/order/9e8d3c44-...
Response:
{
  "maGiaodich": "9e8d3c44-...",
  "status": "ISSUED",
  "paymentUrl": "https://...",
  "policyNumber": "TNDS-...",
  "serialNumber": "AB123456",
  "pdfUrl": "https://piastest.pvi.com.vn/policy/AB123456.pdf"
}''')

add_p('Các trạng thái có thể trả về:')
status_rows = [
    ('status', 'Ý nghĩa'),
    ('SUBMITTING', 'Gateway đang gửi sang PVI, chưa có kết quả'),
    ('SUBMITTED_OK', 'PVI đã nhận, đang chờ cấp GCN (chờ callback hoặc reconcile)'),
    ('SUBMITTED_FAIL', 'PVI từ chối hoặc gateway lỗi khi gửi — xem maGiaodich, có thể tạo lại đơn mới'),
    ('ISSUED', 'Đã cấp GCN — có policyNumber, pdfUrl'),
    ('CALLBACK_TIMEOUT', 'Hết retry mà chưa có kết quả — liên hệ PVIS để xử lý'),
]
t = doc.add_table(rows=len(status_rows), cols=2)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(status_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

# 4.6 Transaction
add_h(2, '4.6. GET /transaction — Danh sách giao dịch của đối tác')
add_p('Trả các giao dịch của chính đối tác (lọc theo Client ID trong header HMAC).')
add_code('''GET /transaction?status=ISSUED&from=2026-06-01&to=2026-06-30

Query params (tất cả optional):
- status:        SUBMITTING | SUBMITTED_OK | SUBMITTED_FAIL | ISSUED | CALLBACK_TIMEOUT
- policyNumber:  số giấy chứng nhận
- from / to:     ngày tạo, ISO format (yyyy-MM-dd)''')

add_h(2, '4.7. GET /transaction/:id — Chi tiết giao dịch')
add_p(
    'Trả chi tiết giao dịch + lịch sử request/response giữa gateway và PVI. '
    'Đối tác chỉ thấy giao dịch của chính mình — gọi ID của partner khác sẽ trả 404.'
)

add_h(2, '4.8. POST /transaction/:id/reconcile — Trigger đối soát thủ công')
add_p(
    'Trong trường hợp đơn ở trạng thái SUBMITTED_OK quá lâu chưa thấy callback, '
    'đối tác có thể trigger đối soát để gateway chủ động hỏi PVI ngay lập tức (không đợi cron).'
)

# ---------- 5. Bảng mã lỗi ----------
add_h(1, '5. Bảng mã lỗi từ PVI Core (passthrough)')
add_p('Trường Status trong response trả về:')
err_rows = [
    ('Status', 'Mô tả'),
    ('00', 'Thực hiện thành công'),
    ('-1', 'Lỗi exception ở PVI Core'),
    ('-105', 'Sai chữ ký (xảy ra ở tầng PVI Core, hiếm khi gặp)'),
    ('-400', 'Dữ liệu sai định dạng'),
    ('-404', 'Dữ liệu không hợp lệ / không tìm thấy'),
]
t = doc.add_table(rows=len(err_rows), cols=2)
t.style = 'Light Grid Accent 1'
for i, row in enumerate(err_rows):
    for j, v in enumerate(row):
        c = t.rows[i].cells[j]
        c.text = ''
        run = c.paragraphs[0].add_run(v)
        if i == 0:
            run.bold = True
            set_cell_bg(c, 'D9E2F3')

# ---------- 6. Quy trình hoàn chỉnh ----------
add_h(1, '6. Quy trình tích hợp đầy đủ (golden path xe máy)')
steps = [
    ('1', 'Lấy danh sách LOẠI XE MÁY', 'POST /api/pvi/catalog với {"ten_dmuc":"LOAIXEMOTOR","parent_value":"1","giatri_chon":""}'),
    ('2', 'Hiển thị form để khách chọn loại xe, biển số, ngày bắt đầu, thời hạn (1/2/3 năm), có mua kèm BH người ngồi trên xe không'),
    ('3', 'Tính ngaycuoi = ngaydau + số_năm ở phía đối tác'),
    ('4', 'Gọi /api/pvi/quote để lấy TotalFee (gửi thamgia_laiphu=true + mtn_laiphu nếu mua kèm)', ''),
    ('5', 'Hiển thị tổng tiền cho khách. Khi khách đồng ý → gọi /api/pvi/order', ''),
    ('6', 'Nhận maGiaodich và paymentUrl. Redirect khách sang paymentUrl', ''),
    ('7', 'Sau khi thanh toán, polling GET /api/pvi/order/:maGiaodich đến khi status = ISSUED', ''),
    ('8', 'Lấy pdfUrl trả cho khách / gửi email', ''),
]
for s in steps:
    p = doc.add_paragraph(style='List Number')
    if len(s) >= 3 and s[2]:
        r1 = p.add_run(s[1])
        r1.bold = True
        p.add_run(' — ')
        p.add_run(s[2])
    else:
        p.add_run(s[1])

# ---------- 7. Hỗ trợ ----------
add_h(1, '7. Hỗ trợ kỹ thuật')
add_p(
    'Khi gặp vấn đề trong quá trình tích hợp, vui lòng liên hệ:'
)
add_kv_table([
    ('Đầu mối kỹ thuật', '_______________________________________'),
    ('Email', '_______________________________________'),
    ('Số điện thoại', '_______________________________________'),
    ('Giờ hỗ trợ', '_______________________________________'),
])
add_p(
    'Khi báo lỗi xin vui lòng đính kèm: Client ID (KHÔNG gửi Secret), maGiaodich (nếu có), '
    'thời điểm xảy ra lỗi, request headers (trừ X-Signature) và response nhận được.'
)

# ---------- Save ----------
import os
out = '/Users/anhtuanbui/Documents/insurance-gateway/docs/Partner-Integration-Guide.docx'
os.makedirs(os.path.dirname(out), exist_ok=True)
doc.save(out)
print('Wrote', out)
